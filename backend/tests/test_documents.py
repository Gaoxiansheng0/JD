import pytest

from resumefit.documents import extract_document


@pytest.fixture
def make_blank_pdf():
    def _make(path):
        from pypdf import PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(width=595, height=842)
        with path.open("wb") as handle:
            writer.write(handle)
        return path

    return _make


def test_extracts_plain_text_with_one_synthetic_page(tmp_path):
    source = tmp_path / "resume.txt"
    source.write_text("项目：智能客服\n负责评测", encoding="utf-8")

    result = extract_document(source, "text/plain")

    assert result.text == "项目：智能客服\n负责评测"
    assert result.pages == 1
    assert result.status == "success"


def test_empty_text_pdf_is_reported_as_scan_not_empty_resume(tmp_path, make_blank_pdf):
    source = make_blank_pdf(tmp_path / "scan.pdf")

    result = extract_document(source, "application/pdf")

    assert result.status == "needs_ocr"
    assert "扫描" in result.warnings[0]


def test_extracts_docx_paragraphs_and_simple_tables(tmp_path):
    from docx import Document

    document = Document()
    document.add_paragraph("项目：智能客服")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "指标"
    table.rows[0].cells[1].text = "首次解决率 71%"
    path = tmp_path / "resume.docx"
    document.save(path)

    result = extract_document(path, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    assert "项目：智能客服" in result.text
    assert "首次解决率 71%" in result.text
    assert result.status == "success"


def test_candidate_facts_are_pending_and_keep_source_offsets(client):
    text = "负责智能客服产品规划。\n建立 200 条测试集，覆盖 12 个场景。\n首次解决率提升至 71%。"
    response = client.post(
        "/api/projects/import",
        json={"name": "智能客服", "text": text, "original_name": "复盘.txt"},
    )

    body = response.json()
    facts = body["facts"]
    assert facts and all(fact["status"] == "pending" for fact in facts)
    for fact in facts:
        assert text[fact["offset_start"] : fact["offset_end"]] == fact["text"]


def test_imported_candidates_are_not_searchable_until_confirmed(client):
    client.post(
        "/api/projects/import",
        json={"name": "智能客服", "text": "建立 200 条测试集，覆盖 12 个场景。", "original_name": "复盘.txt"},
    )

    assert client.get("/api/projects/search", params={"q": "测试集"}).json() == []


def test_master_resume_import_stores_text_without_creating_facts(client):
    client.post(
        "/api/profile/master-resume",
        json={"text": "AI 产品经理\n建立 200 条测试集", "original_name": "简历.txt"},
    )

    assert "测试集" in client.get("/api/profile/master-resume").json()["text"]
    assert client.get("/api/projects/search", params={"q": "测试集"}).json() == []
