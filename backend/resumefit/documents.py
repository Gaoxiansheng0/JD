"""TXT / DOCX / 文本型 PDF 解析与候选事实拆分。

规格 §17：PDF 提取为空判定为扫描件并停止，不返回「空简历」。
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# 句子边界：中文句号/问号/叹号/分号/换行。保留原文偏移，所以用 finditer 而不是 split。
_SENTENCE = re.compile(r"[^。！？；\n]+")
_MIN_FACT_CHARS = 6


@dataclass
class ExtractedDocument:
    text: str
    pages: int
    status: str  # success | needs_ocr | unsupported
    warnings: list[str] = field(default_factory=list)


@dataclass(frozen=True)
class CandidateFact:
    text: str
    offset_start: int
    offset_end: int


def extract_document(file_path: Path, mime_type: str) -> ExtractedDocument:
    suffix = file_path.suffix.lower()

    if mime_type.startswith("text/") or suffix in {".txt", ".md"}:
        return ExtractedDocument(text=file_path.read_text(encoding="utf-8"), pages=1, status="success")

    if mime_type == DOCX_MIME or suffix == ".docx":
        return _extract_docx(file_path)

    if mime_type == "application/pdf" or suffix == ".pdf":
        return _extract_pdf(file_path)

    return ExtractedDocument(
        text="", pages=0, status="unsupported", warnings=[f"不支持的文件类型：{suffix or mime_type}"]
    )


def _extract_docx(file_path: Path) -> ExtractedDocument:
    from docx import Document

    document = Document(str(file_path))
    blocks = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))

    warnings = []
    if document.inline_shapes:
        warnings.append("文档含有图片或图形，其中的文字未被提取")

    return ExtractedDocument(text="\n".join(blocks), pages=1, status="success", warnings=warnings)


def _extract_pdf(file_path: Path) -> ExtractedDocument:
    from pypdf import PdfReader

    reader = PdfReader(str(file_path))
    pages = len(reader.pages)
    text = "\n".join((page.extract_text() or "").strip() for page in reader.pages).strip()

    if not text:
        return ExtractedDocument(
            text="",
            pages=pages,
            status="needs_ocr",
            warnings=["未提取到任何文字，可能是扫描 PDF；请改用截图或直接粘贴文本"],
        )

    return ExtractedDocument(text=text, pages=pages, status="success")


def split_candidate_facts(text: str) -> list[CandidateFact]:
    """把长文本拆成可逐条审阅的候选事实，保留原文偏移。

    只做切分，不做归纳——素材里没有的内容不会凭空出现。
    """
    candidates = []
    for match in _SENTENCE.finditer(text):
        stripped = match.group().strip()
        if len(stripped) < _MIN_FACT_CHARS:
            continue
        start = match.start() + match.group().index(stripped)
        candidates.append(
            CandidateFact(text=stripped, offset_start=start, offset_end=start + len(stripped))
        )
    return candidates
